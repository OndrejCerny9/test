import os
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Inches

from sector_report_agent.app.config import REPORT_VOLUME_PATH


def safe_file_name(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-z0-9_]+", "_", value)
    return value.strip("_")


def generate_docx_report(
    report_name: str,
    sector: str,
    period: str,
    sections: dict,
    chart_paths: Optional[list] = None,
) -> str:
    os.makedirs(REPORT_VOLUME_PATH, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = f"{REPORT_VOLUME_PATH}/{safe_file_name(report_name)}_{timestamp}.docx"

    doc = Document()
    doc.add_heading(report_name, level=0)

    doc.add_paragraph(f"Sector: {sector}")
    doc.add_paragraph(f"Period: {period}")
    doc.add_paragraph("")

    for section_name, section_content in sections.items():
        doc.add_heading(section_name, level=1)
        doc.add_paragraph(section_content or "")

    if chart_paths:
        doc.add_heading("Charts", level=1)
        for chart_path in chart_paths:
            doc.add_paragraph(chart_path)
            doc.add_picture(chart_path, width=Inches(6))

    doc.save(output_path)
    return output_path
