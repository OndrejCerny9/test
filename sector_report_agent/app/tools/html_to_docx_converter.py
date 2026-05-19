"""
HTML to DOCX Converter Tool

Converts HTML report content (including Chart.js configurations) to a Word
document (.docx) with native editable charts and saves it to a Unity Catalog Volume.

Strategy:
- Load corporate DOCX template (with styles, header/footer, page setup)
- Extract Chart.js configs via regex (supports both direct and variable patterns)
- Generate native Word charts (OOXML DrawingML) with embedded Excel data
- Convert HTML text/tables with htmldocx
- Insert native chart objects with captions and source citations
- Insert Table of Contents ("Obsah") after the title
- Replace header placeholders with actual report title/date
- Save to UC Volume via Databricks SDK
"""

import os
import io
import re
import logging
import zipfile
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from htmldocx import HtmlToDocx
from databricks.sdk import WorkspaceClient
from openpyxl import Workbook
from lxml import etree

logger = logging.getLogger(__name__)

DEFAULT_VOLUME_PATH = "/Volumes/agentbricks/volumes/agent_reports"

CHART_PLACEHOLDER = "___CHART_PLACEHOLDER_{}___ "

# Corporate color palette
CORPORATE_COLORS = [
    '4E5B6F',  # Dark blue-gray (primary)
    '007EEA',  # Bright blue
    '898989',  # Medium gray
    'D6ECFF',  # Light blue
    'A7D6FF',  # Sky blue
    'FFDF43',  # Yellow accent
    '8E9BB0',  # Steel blue
    '245375',  # Dark teal
    '4CAF50',  # Green
    'FF6B6B',  # Red accent
]

# Template path
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), "template.docx")

# Chart relationship ID base (use high numbers to avoid conflicts)
CHART_RID_BASE = 100


def get_volume_path() -> str:
    """Get the configured volume path from environment or use default."""
    return os.environ.get("REPORT_VOLUME_PATH", DEFAULT_VOLUME_PATH)


def _find_chart_config_end(text: str) -> int:
    """Find the end of a Chart config object by matching braces."""
    brace_count = 0
    started = False
    for i, ch in enumerate(text):
        if ch == '{':
            brace_count += 1
            started = True
        elif ch == '}':
            brace_count -= 1
            if started and brace_count == 0:
                remaining = text[i:]
                end_match = re.search(r'\)\s*;', remaining)
                if end_match:
                    return i + end_match.end()
                return i + 1
    return len(text)


def _extract_chart_configs(html_content: str) -> list:
    """
    Extract Chart.js configurations from <script> blocks.
    Returns list of dicts with canvas_id and chart config data.
    """
    charts = []
    var_pattern = r"""(?:const|let|var)\s+(\w+)\s*=\s*document\.getElementById\s*\(\s*['\"]([^'\"]+)['\"]\s*\)"""
    var_to_canvas = {}
    for var_match in re.finditer(var_pattern, html_content):
        var_to_canvas[var_match.group(1)] = var_match.group(2)

    pattern1 = r'new\s+Chart\s*\(\s*document\.getElementById\s*\(\s*[\x27\"]([^\x27\"]+)[\x27\"]\s*\)'
    pattern2 = r'new\s+Chart\s*\(\s*(\w+)\s*,'

    found_charts = []
    for match in re.finditer(pattern1, html_content):
        canvas_id = match.group(1)
        rest = html_content[match.end():]
        found_charts.append((canvas_id, rest))

    for match in re.finditer(pattern2, html_content):
        var_name = match.group(1)
        if var_name == 'document':
            continue
        if var_name in var_to_canvas:
            canvas_id = var_to_canvas[var_name]
            if not any(c[0] == canvas_id for c in found_charts):
                rest = html_content[match.end():]
                found_charts.append((canvas_id, rest))

    for canvas_id, rest in found_charts:
        config_end = _find_chart_config_end(rest)
        rest = rest[:config_end]

        type_match = re.search(r'type\s*:\s*[\x27\"]([^\x27\"]+)[\x27\"]', rest[:1000])
        chart_type = type_match.group(1) if type_match else "bar"

        labels_match = re.search(r'labels\s*:\s*\[([^\]]+)\]', rest[:5000])
        labels = []
        if labels_match:
            labels = [l.strip().strip("\x27\\\"") for l in labels_match.group(1).split(",")]

        data_arrays = re.findall(r'data\s*:\s*\[([\d.,\s\-]+)\]', rest[:10000])
        ds_labels = re.findall(r'label\s*:\s*[\x27\"]([^\x27\"]+)[\x27\"]', rest[:10000])
        bg_colors = re.findall(r'backgroundColor\s*:\s*[\x27\"]([^\x27\"]+)[\x27\"]', rest[:10000])
        border_colors = re.findall(r'borderColor\s*:\s*[\x27\"]([^\x27\"]+)[\x27\"]', rest[:10000])

        datasets = []
        for i, data_str in enumerate(data_arrays):
            try:
                data_vals = [float(x.strip()) for x in data_str.split(",") if x.strip()]
                dataset = {"data": data_vals}
                if i < len(ds_labels):
                    dataset["label"] = ds_labels[i]
                if i < len(bg_colors):
                    dataset["backgroundColor"] = bg_colors[i]
                elif i < len(border_colors):
                    dataset["backgroundColor"] = border_colors[i]
                datasets.append(dataset)
            except ValueError:
                continue

        title_text = None
        title_match = re.search(r'text\s*:\s*[\x27\"]([^\x27\"]+)[\x27\"]', rest[:10000])
        if title_match:
            title_text = title_match.group(1)

        if labels and datasets:
            config = {
                "type": chart_type,
                "labels": labels,
                "datasets": datasets,
                "title": title_text,
            }
            charts.append({"canvas_id": canvas_id, "config": config})
            logger.info(f"Extracted chart config for canvas: {canvas_id}")

    return charts


def _parse_color_to_hex(color_str) -> str:
    """Parse rgba/rgb/hex color string to 6-digit hex (no #)."""
    if not isinstance(color_str, str):
        return CORPORATE_COLORS[0]
    color_str = color_str.strip()
    rgba_match = re.match(r'rgba?\(([\d.]+),\s*([\d.]+),\s*([\d.]+)', color_str)
    if rgba_match:
        r = int(float(rgba_match.group(1)))
        g = int(float(rgba_match.group(2)))
        b = int(float(rgba_match.group(3)))
        return f'{r:02X}{g:02X}{b:02X}'
    if color_str.startswith('#'):
        return color_str.lstrip('#').upper()
    return CORPORATE_COLORS[0]


def _create_chart_xml(config: dict, chart_index: int) -> bytes:
    """
    Create a native Word chart XML (DrawingML c:chartSpace) from chart config.
    Supports bar, line, pie, doughnut.
    """
    chart_type = config.get("type", "bar")
    labels = config.get("labels", [])
    datasets = config.get("datasets", [])
    title = config.get("title")

    C = 'http://schemas.openxmlformats.org/drawingml/2006/chart'
    A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    nsmap = {'c': C, 'a': A, 'r': R}

    chartSpace = etree.Element(f'{{{C}}}chartSpace', nsmap=nsmap)
    chart_el = etree.SubElement(chartSpace, f'{{{C}}}chart')

    # Title
    if title:
        title_el = etree.SubElement(chart_el, f'{{{C}}}title')
        tx = etree.SubElement(title_el, f'{{{C}}}tx')
        rich = etree.SubElement(tx, f'{{{C}}}rich')
        etree.SubElement(rich, f'{{{A}}}bodyPr')
        etree.SubElement(rich, f'{{{A}}}lstStyle')
        p = etree.SubElement(rich, f'{{{A}}}p')
        pPr = etree.SubElement(p, f'{{{A}}}pPr')
        defRPr = etree.SubElement(pPr, f'{{{A}}}defRPr')
        defRPr.set('sz', '1100'); defRPr.set('b', '1')
        r = etree.SubElement(p, f'{{{A}}}r')
        rPr = etree.SubElement(r, f'{{{A}}}rPr')
        rPr.set('lang', 'cs-CZ'); rPr.set('sz', '1100'); rPr.set('b', '1')
        t_el = etree.SubElement(r, f'{{{A}}}t')
        t_el.text = title
        etree.SubElement(title_el, f'{{{C}}}overlay').set('val', '0')

    etree.SubElement(chart_el, f'{{{C}}}autoTitleDeleted').set('val', '0')
    plotArea = etree.SubElement(chart_el, f'{{{C}}}plotArea')
    etree.SubElement(plotArea, f'{{{C}}}layout')

    # Determine chart element type
    if chart_type in ('pie', 'doughnut'):
        chart_elem_tag = f'{{{C}}}pieChart' if chart_type == 'pie' else f'{{{C}}}doughnutChart'
        chart_elem = etree.SubElement(plotArea, chart_elem_tag)
        etree.SubElement(chart_elem, f'{{{C}}}varyColors').set('val', '1')

        # For pie/doughnut, use first dataset only
        if datasets:
            ds = datasets[0]
            ser = etree.SubElement(chart_elem, f'{{{C}}}ser')
            etree.SubElement(ser, f'{{{C}}}idx').set('val', '0')
            etree.SubElement(ser, f'{{{C}}}order').set('val', '0')

            # Series name
            tx = etree.SubElement(ser, f'{{{C}}}tx')
            strRef = etree.SubElement(tx, f'{{{C}}}strRef')
            etree.SubElement(strRef, f'{{{C}}}f').text = "Sheet1!$B$1"
            sc = etree.SubElement(strRef, f'{{{C}}}strCache')
            etree.SubElement(sc, f'{{{C}}}ptCount').set('val', '1')
            pt = etree.SubElement(sc, f'{{{C}}}pt'); pt.set('idx', '0')
            etree.SubElement(pt, f'{{{C}}}v').text = ds.get('label', 'Data')

            # Color each slice
            for i in range(len(labels)):
                dPt = etree.SubElement(ser, f'{{{C}}}dPt')
                etree.SubElement(dPt, f'{{{C}}}idx').set('val', str(i))
                spPr = etree.SubElement(dPt, f'{{{C}}}spPr')
                sf = etree.SubElement(spPr, f'{{{A}}}solidFill')
                etree.SubElement(sf, f'{{{A}}}srgbClr').set('val', CORPORATE_COLORS[i % len(CORPORATE_COLORS)])

            # Categories
            cat_el = etree.SubElement(ser, f'{{{C}}}cat')
            sr_c = etree.SubElement(cat_el, f'{{{C}}}strRef')
            etree.SubElement(sr_c, f'{{{C}}}f').text = f"Sheet1!$A$2:$A${len(labels)+1}"
            sc_c = etree.SubElement(sr_c, f'{{{C}}}strCache')
            etree.SubElement(sc_c, f'{{{C}}}ptCount').set('val', str(len(labels)))
            for i, cn in enumerate(labels):
                pc = etree.SubElement(sc_c, f'{{{C}}}pt'); pc.set('idx', str(i))
                etree.SubElement(pc, f'{{{C}}}v').text = cn

            # Values
            val_el = etree.SubElement(ser, f'{{{C}}}val')
            nr = etree.SubElement(val_el, f'{{{C}}}numRef')
            etree.SubElement(nr, f'{{{C}}}f').text = f"Sheet1!$B$2:$B${len(ds['data'])+1}"
            nc = etree.SubElement(nr, f'{{{C}}}numCache')
            etree.SubElement(nc, f'{{{C}}}formatCode').text = '#,##0'
            etree.SubElement(nc, f'{{{C}}}ptCount').set('val', str(len(ds['data'])))
            for i, v in enumerate(ds['data']):
                pv = etree.SubElement(nc, f'{{{C}}}pt'); pv.set('idx', str(i))
                etree.SubElement(pv, f'{{{C}}}v').text = str(v)

        if chart_type == 'doughnut':
            etree.SubElement(chart_elem, f'{{{C}}}holeSize').set('val', '50')

    else:
        # Bar or Line chart
        if chart_type == 'line':
            chart_elem = etree.SubElement(plotArea, f'{{{C}}}lineChart')
            etree.SubElement(chart_elem, f'{{{C}}}grouping').set('val', 'standard')
        else:
            chart_elem = etree.SubElement(plotArea, f'{{{C}}}barChart')
            etree.SubElement(chart_elem, f'{{{C}}}barDir').set('val', 'col')
            etree.SubElement(chart_elem, f'{{{C}}}grouping').set('val', 'clustered')

        etree.SubElement(chart_elem, f'{{{C}}}varyColors').set('val', '0')

        for idx, ds in enumerate(datasets):
            ser = etree.SubElement(chart_elem, f'{{{C}}}ser')
            etree.SubElement(ser, f'{{{C}}}idx').set('val', str(idx))
            etree.SubElement(ser, f'{{{C}}}order').set('val', str(idx))

            # Series name
            tx = etree.SubElement(ser, f'{{{C}}}tx')
            strRef = etree.SubElement(tx, f'{{{C}}}strRef')
            etree.SubElement(strRef, f'{{{C}}}f').text = f"Sheet1!${chr(66+idx)}$1"
            sc = etree.SubElement(strRef, f'{{{C}}}strCache')
            etree.SubElement(sc, f'{{{C}}}ptCount').set('val', '1')
            pt = etree.SubElement(sc, f'{{{C}}}pt'); pt.set('idx', '0')
            etree.SubElement(pt, f'{{{C}}}v').text = ds.get('label', f'Series {idx+1}')

            # Color
            color_hex = CORPORATE_COLORS[idx % len(CORPORATE_COLORS)]
            if 'backgroundColor' in ds:
                color_hex = _parse_color_to_hex(ds['backgroundColor'])
            spPr = etree.SubElement(ser, f'{{{C}}}spPr')
            sf = etree.SubElement(spPr, f'{{{A}}}solidFill')
            etree.SubElement(sf, f'{{{A}}}srgbClr').set('val', color_hex)

            if chart_type == 'line':
                ln = etree.SubElement(spPr, f'{{{A}}}ln')
                ln.set('w', '28575')
                sf_ln = etree.SubElement(ln, f'{{{A}}}solidFill')
                etree.SubElement(sf_ln, f'{{{A}}}srgbClr').set('val', color_hex)
                # Marker
                marker = etree.SubElement(ser, f'{{{C}}}marker')
                etree.SubElement(marker, f'{{{C}}}symbol').set('val', 'circle')
                etree.SubElement(marker, f'{{{C}}}size').set('val', '5')

            # Categories
            cat_el = etree.SubElement(ser, f'{{{C}}}cat')
            sr_c = etree.SubElement(cat_el, f'{{{C}}}strRef')
            etree.SubElement(sr_c, f'{{{C}}}f').text = f"Sheet1!$A$2:$A${len(labels)+1}"
            sc_c = etree.SubElement(sr_c, f'{{{C}}}strCache')
            etree.SubElement(sc_c, f'{{{C}}}ptCount').set('val', str(len(labels)))
            for i, cn in enumerate(labels):
                pc = etree.SubElement(sc_c, f'{{{C}}}pt'); pc.set('idx', str(i))
                etree.SubElement(pc, f'{{{C}}}v').text = cn

            # Values
            val_el = etree.SubElement(ser, f'{{{C}}}val')
            numRef = etree.SubElement(val_el, f'{{{C}}}numRef')
            etree.SubElement(numRef, f'{{{C}}}f').text = f"Sheet1!${chr(66+idx)}$2:${chr(66+idx)}${len(ds['data'])+1}"
            nc = etree.SubElement(numRef, f'{{{C}}}numCache')
            etree.SubElement(nc, f'{{{C}}}formatCode').text = '#,##0'
            etree.SubElement(nc, f'{{{C}}}ptCount').set('val', str(len(ds['data'])))
            for i, v in enumerate(ds['data']):
                pv = etree.SubElement(nc, f'{{{C}}}pt'); pv.set('idx', str(i))
                etree.SubElement(pv, f'{{{C}}}v').text = str(v)

        # Axes (not for pie/doughnut)
        etree.SubElement(chart_elem, f'{{{C}}}axId').set('val', '1')
        etree.SubElement(chart_elem, f'{{{C}}}axId').set('val', '2')

        catAx = etree.SubElement(plotArea, f'{{{C}}}catAx')
        etree.SubElement(catAx, f'{{{C}}}axId').set('val', '1')
        s1 = etree.SubElement(catAx, f'{{{C}}}scaling')
        etree.SubElement(s1, f'{{{C}}}orientation').set('val', 'minMax')
        etree.SubElement(catAx, f'{{{C}}}delete').set('val', '0')
        etree.SubElement(catAx, f'{{{C}}}axPos').set('val', 'b')
        etree.SubElement(catAx, f'{{{C}}}crossAx').set('val', '2')

        valAx = etree.SubElement(plotArea, f'{{{C}}}valAx')
        etree.SubElement(valAx, f'{{{C}}}axId').set('val', '2')
        s2 = etree.SubElement(valAx, f'{{{C}}}scaling')
        etree.SubElement(s2, f'{{{C}}}orientation').set('val', 'minMax')
        etree.SubElement(valAx, f'{{{C}}}delete').set('val', '0')
        etree.SubElement(valAx, f'{{{C}}}axPos').set('val', 'l')
        etree.SubElement(valAx, f'{{{C}}}crossAx').set('val', '1')

    # Legend
    legend = etree.SubElement(chart_el, f'{{{C}}}legend')
    etree.SubElement(legend, f'{{{C}}}legendPos').set('val', 'b')
    etree.SubElement(legend, f'{{{C}}}overlay').set('val', '0')
    etree.SubElement(chart_el, f'{{{C}}}plotVisOnly').set('val', '1')

    # External data reference
    extData = etree.SubElement(chartSpace, f'{{{C}}}externalData')
    extData.set(f'{{{R}}}id', 'rId1')
    etree.SubElement(extData, f'{{{C}}}autoUpdate').set('val', '0')

    return etree.tostring(chartSpace, xml_declaration=True, encoding='UTF-8', standalone=True)


def _create_chart_xlsx(config: dict) -> bytes:
    """Create embedded Excel workbook with chart data."""
    labels = config.get("labels", [])
    datasets = config.get("datasets", [])

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.cell(1, 1, "Category")
    for i, ds in enumerate(datasets):
        ws.cell(1, i + 2, ds.get('label', f'Series {i+1}'))
    for row_idx, label in enumerate(labels):
        ws.cell(row_idx + 2, 1, label)
        for col_idx, ds in enumerate(datasets):
            if row_idx < len(ds.get('data', [])):
                ws.cell(row_idx + 2, col_idx + 2, ds['data'][row_idx])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _prepare_html_and_charts(html_content: str) -> tuple:
    """
    Process HTML: extract charts, replace canvas elements with placeholders,
    remove scripts/styles, return clean HTML and chart configs.
    """
    charts = _extract_chart_configs(html_content)
    chart_data = {}

    for chart_info in charts:
        canvas_id = chart_info["canvas_id"]
        config = chart_info["config"]
        chart_data[canvas_id] = {"config": config, "title": config.get("title")}
        logger.info(f"Prepared chart config: {canvas_id}")

    soup = BeautifulSoup(html_content, "html.parser")
    for canvas_id in chart_data:
        canvas = soup.find("canvas", {"id": canvas_id})
        if canvas:
            placeholder = soup.new_tag("p")
            placeholder.string = CHART_PLACEHOLDER.format(canvas_id)
            parent = canvas.parent
            if parent and parent.name == "div":
                parent.replace_with(placeholder)
            else:
                canvas.replace_with(placeholder)

    for tag in soup.find_all(["script", "style"]):
        tag.decompose()

    body = soup.find("body")
    clean_html = str(body) if body else str(soup)
    return clean_html, chart_data


def _load_template() -> Document:
    """Load the corporate DOCX template."""
    if os.path.exists(TEMPLATE_PATH):
        logger.info(f"Loading template from: {TEMPLATE_PATH}")
        return Document(TEMPLATE_PATH)
    else:
        logger.warning(f"Template not found at {TEMPLATE_PATH}, using blank document")
        return Document()


def _replace_header_placeholders(doc: Document, title: str, date_str: str):
    """Replace {{REPORT_TITLE}} and {{REPORT_DATE}} in the header."""
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                if "{{REPORT_TITLE}}" in run.text:
                    run.text = run.text.replace("{{REPORT_TITLE}}", title)
                if "{{REPORT_DATE}}" in run.text:
                    run.text = run.text.replace("{{REPORT_DATE}}", date_str)
        footer = section.footer
        for paragraph in footer.paragraphs:
            for run in paragraph.runs:
                if "{{COMPANY_NAME}}" in run.text:
                    run.text = run.text.replace("{{COMPANY_NAME}}", "")


def _extract_report_title(html_content: str) -> str:
    """Extract the report title from the HTML (first h1 tag)."""
    soup = BeautifulSoup(html_content, "html.parser")
    h1 = soup.find("h1")
    if h1:
        return h1.get_text(strip=True)
    return "Sector Report"


def _insert_table_of_contents(doc: Document):
    """
    Insert a Table of Contents page after the first heading (title).
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    insert_after_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.style and 'Heading 1' in para.style.name:
            insert_after_idx = i
            break

    if insert_after_idx is None:
        insert_after_idx = 0

    title_element = doc.paragraphs[insert_after_idx]._element

    # "Obsah" heading
    obsah_p = OxmlElement('w:p')
    obsah_pPr = OxmlElement('w:pPr')
    obsah_pStyle = OxmlElement('w:pStyle')
    obsah_pStyle.set(_qn('w:val'), 'Heading1')
    obsah_pPr.append(obsah_pStyle)
    obsah_p.append(obsah_pPr)
    obsah_r = OxmlElement('w:r')
    obsah_t = OxmlElement('w:t')
    obsah_t.text = 'Obsah'
    obsah_r.append(obsah_t)
    obsah_p.append(obsah_r)

    # TOC field
    toc_p = OxmlElement('w:p')
    r_begin = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(_qn('w:fldCharType'), 'begin')
    r_begin.append(fldChar_begin)
    toc_p.append(r_begin)

    r_instr = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(_qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "2-3" \\h \\z \\u '
    r_instr.append(instrText)
    toc_p.append(r_instr)

    r_sep = OxmlElement('w:r')
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(_qn('w:fldCharType'), 'separate')
    r_sep.append(fldChar_sep)
    toc_p.append(r_sep)

    r_placeholder = OxmlElement('w:r')
    t_placeholder = OxmlElement('w:t')
    t_placeholder.text = 'Right-click and select "Update Field" to populate'
    r_placeholder.append(t_placeholder)
    toc_p.append(r_placeholder)

    r_end = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(_qn('w:fldCharType'), 'end')
    r_end.append(fldChar_end)
    toc_p.append(r_end)

    # Page break
    pagebreak_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(_qn('w:type'), 'page')
    pb_r.append(pb_br)
    pagebreak_p.append(pb_r)

    # Insert in reverse order
    title_element.addnext(pagebreak_p)
    title_element.addnext(toc_p)
    title_element.addnext(obsah_p)
    logger.info("Inserted Table of Contents (Obsah) after title")


def _apply_corporate_styling(doc: Document):
    """Post-process the document to apply corporate fonts."""
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else "Normal"
        if 'Heading' in style_name:
            target_font = 'Franklin Gothic Book'
            target_color = RGBColor(0x24, 0x53, 0x75)
        elif 'Chart Title' in style_name:
            target_font = 'Times New Roman'
            target_color = None
        elif 'Chart Source' in style_name:
            target_font = 'Times New Roman'
            target_color = RGBColor(0x20, 0x20, 0x20)
        else:
            target_font = 'Franklin Gothic Book'
            target_color = None

        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = target_font
            if target_color and 'Heading' in style_name:
                if not run.font.color.rgb:
                    run.font.color.rgb = target_color

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not run.font.name:
                            run.font.name = 'Franklin Gothic Book'
                        if not run.font.size:
                            run.font.size = Pt(9)


def _build_docx_with_charts(html_content: str, chart_data: dict, report_title: str) -> Document:
    """
    Build DOCX using the corporate template with native Word charts.
    Two-pass approach:
    1. Build document with placeholder paragraphs containing chart drawing references
    2. Post-process the ZIP to inject chart XML and embedded Excel files
    """
    doc = _load_template()
    date_str = datetime.now().strftime("%B %Y")
    _replace_header_placeholders(doc, report_title, date_str)

    # Convert HTML
    parser = HtmlToDocx()
    parser.add_html_to_document(html_content, doc)

    # Insert TOC
    _insert_table_of_contents(doc)

    # Track chart insertion info
    chart_insertions = []  # list of (chart_index, canvas_id, config)

    # Find and replace placeholder paragraphs with chart drawing references
    chart_index = 0
    for canvas_id, data in chart_data.items():
        placeholder_text = CHART_PLACEHOLDER.format(canvas_id).strip()
        config = data["config"]
        chart_title = data.get("title")

        for i, paragraph in enumerate(doc.paragraphs):
            if placeholder_text in paragraph.text:
                paragraph.clear()

                # Add chart title above
                if chart_title:
                    title_para = paragraph.insert_paragraph_before(chart_title)
                    try:
                        title_para.style = doc.styles['Chart Title']
                    except KeyError:
                        for run in title_para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(10)
                            run.font.bold = True

                # Insert inline chart drawing reference
                chart_rid = f'rId{CHART_RID_BASE + chart_index}'
                cx = int(16 * 914400 / 2.54)  # 16cm width
                cy = int(9 * 914400 / 2.54)   # 9cm height

                W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                WP = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                A = 'http://schemas.openxmlformats.org/drawingml/2006/main'
                R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
                C_NS = 'http://schemas.openxmlformats.org/drawingml/2006/chart'

                run = paragraph.add_run()
                drawing = OxmlElement('w:drawing')

                inline = etree.SubElement(drawing, f'{{{WP}}}inline')
                inline.set('distT', '0'); inline.set('distB', '0')
                inline.set('distL', '0'); inline.set('distR', '0')
                ext_el = etree.SubElement(inline, f'{{{WP}}}extent')
                ext_el.set('cx', str(cx)); ext_el.set('cy', str(cy))
                eff = etree.SubElement(inline, f'{{{WP}}}effectExtent')
                eff.set('l', '0'); eff.set('t', '0'); eff.set('r', '0'); eff.set('b', '0')
                dp = etree.SubElement(inline, f'{{{WP}}}docPr')
                dp.set('id', str(chart_index + 10)); dp.set('name', f'Chart {chart_index + 1}')
                etree.SubElement(inline, f'{{{WP}}}cNvGraphicFramePr')
                graphic = etree.SubElement(inline, f'{{{A}}}graphic')
                gd = etree.SubElement(graphic, f'{{{A}}}graphicData')
                gd.set('uri', C_NS)
                cr = etree.SubElement(gd, f'{{{C_NS}}}chart')
                cr.set(f'{{{R_NS}}}id', chart_rid)

                run._r.append(drawing)

                # Source citation after chart
                new_p_elem = OxmlElement('w:p')
                new_r_elem = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
                rPr.append(sz)
                italic_el = OxmlElement('w:i'); rPr.append(italic_el)
                color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '202020')
                rPr.append(color_el)
                new_r_elem.append(rPr)
                t_elem = OxmlElement('w:t')
                t_elem.text = "Zdroj: Vlastní zpracování"
                new_r_elem.append(t_elem)
                new_p_elem.append(new_r_elem)
                paragraph._element.addnext(new_p_elem)

                chart_insertions.append((chart_index, canvas_id, config))
                chart_index += 1
                logger.info(f"Inserted native chart reference: {canvas_id} (chart{chart_index})")
                break

    _apply_corporate_styling(doc)
    return doc, chart_insertions


def _inject_charts_into_docx(doc_bytes: bytes, chart_insertions: list) -> bytes:
    """
    Post-process the .docx ZIP to inject chart XML parts and embedded Excel files.
    """
    if not chart_insertions:
        return doc_bytes

    input_buf = io.BytesIO(doc_bytes)
    output_buf = io.BytesIO()

    with zipfile.ZipFile(input_buf, 'r') as zin:
        with zipfile.ZipFile(output_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == '[Content_Types].xml':
                    tree = etree.fromstring(data)
                    ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
                    for ci, _, _ in chart_insertions:
                        ov = etree.SubElement(tree, f'{{{ns}}}Override')
                        ov.set('PartName', f'/word/charts/chart{ci+1}.xml')
                        ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.drawingml.chart+xml')
                        ov2 = etree.SubElement(tree, f'{{{ns}}}Override')
                        ov2.set('PartName', f'/word/embeddings/Microsoft_Excel_Worksheet{ci+1}.xlsx')
                        ov2.set('ContentType', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif item.filename == 'word/_rels/document.xml.rels':
                    tree = etree.fromstring(data)
                    ns_r = 'http://schemas.openxmlformats.org/package/2006/relationships'
                    for ci, _, _ in chart_insertions:
                        nr = etree.SubElement(tree, f'{{{ns_r}}}Relationship')
                        nr.set('Id', f'rId{CHART_RID_BASE + ci}')
                        nr.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/chart')
                        nr.set('Target', f'charts/chart{ci+1}.xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

            # Add chart files
            for ci, canvas_id, config in chart_insertions:
                chart_xml = _create_chart_xml(config, ci)
                xlsx_data = _create_chart_xlsx(config)
                zout.writestr(f'word/charts/chart{ci+1}.xml', chart_xml)
                zout.writestr(f'word/embeddings/Microsoft_Excel_Worksheet{ci+1}.xlsx', xlsx_data)
                # Chart rels
                chart_rels = (
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                    f'<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/package" Target="../embeddings/Microsoft_Excel_Worksheet{ci+1}.xlsx"/>'
                    '</Relationships>'
                )
                zout.writestr(f'word/charts/_rels/chart{ci+1}.xml.rels', chart_rels)

    output_buf.seek(0)
    return output_buf.getvalue()


def _save_to_volume(doc_bytes: bytes, volume_path: str, filename: str) -> str:
    """Save DOCX bytes to a UC Volume using the Databricks SDK."""
    full_path = f"{volume_path}/{filename}"
    w = WorkspaceClient()
    w.files.upload(full_path, io.BytesIO(doc_bytes), overwrite=True)
    return full_path


def convert_html_to_docx(filename: str, html_content: str) -> dict:
    """
    Convert HTML content (including Chart.js) to a styled Word document
    with native editable charts, and save to the configured UC Volume.
    """
    volume_path = get_volume_path()
    logger.info(f"Target volume path: {volume_path}")

    try:
        report_title = _extract_report_title(html_content)
        soup = BeautifulSoup(html_content, "html.parser")
        has_scripts = bool(soup.find_all("script"))

        if has_scripts:
            logger.info("HTML contains scripts - generating native Word charts")
            clean_html, chart_data = _prepare_html_and_charts(html_content)
            doc, chart_insertions = _build_docx_with_charts(clean_html, chart_data, report_title)

            # Save doc to bytes, then inject chart parts
            buf = io.BytesIO()
            doc.save(buf)
            doc_bytes = buf.getvalue()
            doc_bytes = _inject_charts_into_docx(doc_bytes, chart_insertions)
        else:
            logger.info("HTML is static - converting directly with template")
            body = soup.find("body")
            clean_html = str(body) if body else html_content
            doc = _load_template()
            date_str = datetime.now().strftime("%B %Y")
            _replace_header_placeholders(doc, report_title, date_str)
            parser_obj = HtmlToDocx()
            parser_obj.add_html_to_document(clean_html, doc)
            _insert_table_of_contents(doc)
            _apply_corporate_styling(doc)
            buf = io.BytesIO()
            doc.save(buf)
            doc_bytes = buf.getvalue()

    except Exception as e:
        logger.error(f"HTML to DOCX conversion failed: {e}")
        raise ValueError(f"HTML to DOCX conversion failed: {e}") from e

    try:
        full_path = _save_to_volume(doc_bytes, volume_path, filename)
        logger.info(f"DOCX report saved successfully: {full_path}")
    except Exception as e:
        logger.error(f"Failed to write report to volume: {e}")
        raise OSError(f"Failed to write report to {volume_path}/{filename}: {e}") from e

    return {
        "status": "success",
        "path": full_path,
        "filename": filename,
    }
